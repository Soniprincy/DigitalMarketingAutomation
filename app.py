import streamlit as st
import os
import textwrap
from PIL import Image, ImageDraw, ImageFont
from moviepy.editor import ImageClip
from docx import Document
import shutil  # for zipping
import io

# ================================
# FUNCTIONS
# ================================

def create_banner(text, filename, color=(25, 90, 200), outdir="output"):
    os.makedirs(outdir, exist_ok=True)
    img = Image.new('RGB', (1200, 600), color=color)
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 50)
    except:
        font = ImageFont.load_default()

    lines = textwrap.wrap(text, width=20)
    y_text = 200
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text(((1200 - w) / 2, y_text), line, fill="white", font=font)
        y_text += h + 10

    img.save(os.path.join(outdir, filename + ".jpg"))
    img.save(os.path.join(outdir, filename + ".png"))


def create_text_files(topic, outdir):
    tagline = f"Empower yourself with {topic} â€” from basics to mastery!"
    description = (
        f"The {topic} is designed to introduce learners to the core concepts of data science. "
        f"From understanding data to building models, this program helps everyone get started "
        f"with real-world insights. Learn, explore, and transform your future!"
    )
    keywords = "Data Science, Machine Learning, AI, Python, Training, Basics, Beginners, Career"

    with open(os.path.join(outdir, "tagline.txt"), "w", encoding="utf-8") as f:
        f.write(tagline)
    with open(os.path.join(outdir, "description.txt"), "w", encoding="utf-8") as f:
        f.write(description)
    with open(os.path.join(outdir, "keywords.txt"), "w", encoding="utf-8") as f:
        f.write(keywords)


def create_video(topic, outdir):
    os.makedirs(outdir, exist_ok=True)
    img = Image.new('RGB', (1280, 720), color=(30, 30, 90))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 60)
    except:
        font = ImageFont.load_default()

    text = topic
    bbox = draw.textbbox((0, 0), text, font=font)
    w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((1280 - w) / 2, (720 - h) / 2), text, fill="white", font=font)

    img_path = os.path.join(outdir, "video_slide.png")
    img.save(img_path)

    clip = ImageClip(img_path).set_duration(10)  # shorter duration for demo
    clip.write_videofile(os.path.join(outdir, "short_video.mp4"), fps=24, verbose=False, logger=None)


def create_blog(topic, outdir):
    blog = Document()
    blog.add_heading(f"{topic}", level=1)
    blog.add_paragraph(
        f"The {topic} introduces you to key data science concepts â€” "
        "from data collection to visualization and machine learning. "
        "A complete beginner-friendly roadmap to mastering data science."
    )
    blog.add_paragraph("\nKey Learning Points:")
    blog.add_paragraph("â€¢ Introduction to Data Science Concepts")
    blog.add_paragraph("â€¢ Data Cleaning and Exploration")
    blog.add_paragraph("â€¢ Visualization Techniques")
    blog.add_paragraph("â€¢ Simple Machine Learning Models")
    blog.add_paragraph("\nEnroll today to start your Data Science journey!")
    blog.save(os.path.join(outdir, "blog.docx"))


def create_linkedin_post(topic, link, outdir):
    post = f"""
ğŸš€ Excited to share our latest launch â€” {topic}! ğŸ¯

ğŸ“Š Learn how to collect, clean, and analyze data.
ğŸ’¡ Build a strong foundation in data science and machine learning.
ğŸ“ Perfect for beginners looking to start their data career.

Register now ğŸ‘‰ {link}
#DataScience #AI #Learning #Education #CareerGrowth
"""
    with open(os.path.join(outdir, "linkedin_post.txt"), "w", encoding="utf-8") as f:
        f.write(post)


def create_offer_banners(topic, outdir):
    offer_texts = [
        "ğŸ‰ Flat 50% Off â€“ Enroll Today!",
        "ğŸ”¥ Early Bird Offer â€“ Limited Seats!",
        "ğŸš€ Learn Data Science Now â€“ Pay Later!",
        "ğŸ’» Master Data Science in 30 Days!"
    ]
    colors = [(200, 50, 50), (50, 150, 50), (50, 80, 200), (220, 140, 0)]
    for i, (text, color) in enumerate(zip(offer_texts, colors)):
        create_banner(f"{topic}\n{text}", f"offer_banner_{i+1}", color=color, outdir=outdir)


def generate_all(topic, link):
    outdir = "output"
    os.makedirs(outdir, exist_ok=True)

    create_banner(topic, "main_banner", color=(25, 90, 200), outdir=outdir)
    create_text_files(topic, outdir)
    create_video(topic, outdir)
    create_blog(topic, outdir)
    create_linkedin_post(topic, link, outdir)
    create_offer_banners(topic, outdir)

    # Create ZIP file in memory
    zip_buffer = io.BytesIO()
    shutil.make_archive("marketing_assets", 'zip', outdir)
    with open("marketing_assets.zip", "rb") as f:
        zip_buffer.write(f.read())

    zip_buffer.seek(0)
    return zip_buffer


# ================================
# STREAMLIT UI
# ================================
st.set_page_config(page_title="Digital Marketing Automation", page_icon="ğŸ¯", layout="centered")

st.title("ğŸ¯ Digital Marketing Automation Tool")
st.write("Automatically generate banners, videos, blogs, LinkedIn posts, and offer posters for your marketing campaign.")

topic = st.text_input("Enter Topic:", "Data Science Basic Training Program for Everyone")
link = st.text_input("Enter Link:", "https://example.com")

if st.button("ğŸš€ Generate & Download Files"):
    with st.spinner("âœ¨ Generating all marketing materials... please wait."):
        zip_data = generate_all(topic, link)
        st.success("âœ… All files generated successfully!")

        st.download_button(
            label="â¬‡ï¸ Download All Marketing Files (ZIP)",
            data=zip_data,
            file_name="marketing_assets.zip",
            mime="application/zip"
        )

st.markdown("---")
st.caption("All generated files are packaged as a downloadable ZIP.")
